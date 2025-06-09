import os
import pandas as pd
from datetime import datetime

from docx.api import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentObject
from docx.parts.document import DocumentPart


class DocxParser:
   """Парсер DOCX-файлов для извлечения данных."""

   def load_document(self, path_to_docx):
      return Document(docx=path_to_docx)

   def get_table_as_lists(
         self, doc: DocumentObject,
         table_index: int=0
      ) -> list[list[str | int | float]]:
      table_datalist = list()
      table = doc.tables[table_index]
      for row in table.rows:
         row_datalist = list()
         for cell in row.cells:
            text = cell.text.strip()
            row_datalist.append(text)
         table_datalist.append(row_datalist)
      return table_datalist

   def get_paragraph(self, doc: DocumentObject, p_idx: int = 0) -> str:
      paragraph = doc.paragraphs[p_idx]
      if paragraph.text.strip():
         return paragraph.text

   def transform_to_dataframe(self, datalist: list[list], headers_count: int = 1) -> pd.DataFrame:
      common = [a == b for a, b in zip(datalist[0], datalist[1])]
      header_rows = datalist[:2]
      body_data = datalist[2:]
      columns = pd.MultiIndex.from_arrays(header_rows)
      df = pd.DataFrame(body_data, columns=columns)
      return df

   @staticmethod
   def _transform_to_dataframe(datalist: list[list], headers_count: int = 1) -> pd.DataFrame:
      if headers_count == 1:
         headers = datalist[0]
         data = datalist[1:]
      elif headers_count > 1:
         raw_headers = datalist[:headers_count]
         merge_headers = dict()
         for idx, header_list in enumerate(raw_headers):
            columns_count = len(header_list)
            for col_index in range(columns_count):
               if header_list[col_index] == header_list[col_index+1]:
                  merge_headers[col_index] = header_list[col_index]
         data = datalist[headers_count:]
      return pd.DataFrame(data, columns=headers)

   def _process_cell(self):
      """Метод конвертации значений в ячейках таблицы"""
      pass

   def iter_content(self):
      for i, element in enumerate(self.doc.iter_inner_content()):
         if isinstance(element, Paragraph):
            pass
         elif isinstance(element, Table):
            pass

   def get_table_to_index(self):
      pass

   def get_all_tables(self):
      pass

   def convert_table_to_dict(self):
      pass

   @staticmethod
   def compute_for_output_table_4(dfs_list: list[pd.DataFrame]):
      df_1 = dfs_list[0]
      df_2 = dfs_list[1]
      df_3 = dfs_list[2]

      now = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
      data = {
        "Наименование продукции": df_1.iloc[:, 0].values,
        "Количество на складе": (df_2.iloc[:, 1].astype(int).values +
                                 df_3.iloc[:, 1].astype(int).values -
                                 df_1.iloc[:, 1].astype(int).values),
        "Единица измерения": df_1.iloc[:, 2].values,
        "Дата": [now] * len(df_1)
      }

      output_df = pd.DataFrame(data, columns=data.keys())
      n_cols = len(output_df.columns)
      numeration_row = list(range(1, n_cols + 1))
      output_df.loc[-1] = numeration_row
      output_df = output_df.sort_index().reset_index(drop=True)
      output_df = output_df.astype(str)
      return output_df
