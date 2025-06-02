import os
import docx
import pandas as pd
from datetime import datetime


class DocxParser:
   def __init__(self, doc_path: str=None):
      self.doc = self.load_document(doc_path)

   @staticmethod
   def load_document(path_to_docx):
      return docx.Document(docx=path_to_docx)

   def transform_to_list_of_lists(self) -> list[list[str | int | float]]:
      table_datalist = list()
      for table in self.doc.tables:
         for row in table.rows:
            row_datalist = list()
            for cell in row.cells:
               row_datalist.append(cell.text)
            table_datalist.append(row_datalist)
      return table_datalist

   @staticmethod
   def transform_to_dataframe(datalist: list[list], headers_count: int = 1) -> pd.DataFrame:
      if headers_count == 1:
         headers = datalist[0]
         data = datalist[1:]
      elif headers_count > 1:
         raw_headers = datalist[:headers_count]
         merge_headers = dict()
         for idx, header_list in enumerate(raw_headers):
            columns_count = len(header_list)
            for col_index in range(columns_count):
               if header_list[col_index] == header_list[col_index+1]:
                  merge_headers[col_index] = header_list[col_index]
         print(merge_headers)
         data = datalist[headers_count:]
      return pd.DataFrame(data, columns=headers)

   def get_table_name(self):
      for paragraph in self.doc.paragraphs:
         if paragraph.text.strip():
            return paragraph.text

   @staticmethod
   def compute_for_output_table_4(dfs_list: list[pd.DataFrame]):
      df_1 = dfs_list[0]
      df_2 = dfs_list[1]
      df_3 = dfs_list[2]

      now = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
      data = {
        "Наименование продукции": df_1.iloc[:, 0].values,
        "Количество на складе": (df_2.iloc[:, 1].astype(int).values +
                                 df_3.iloc[:, 1].astype(int).values -
                                 df_1.iloc[:, 1].astype(int).values),
        "Единица измерения": df_1.iloc[:, 2].values,
        "Дата": [now] * len(df_1)
      }

      output_df = pd.DataFrame(data, columns=data.keys())
      n_cols = len(output_df.columns)
      numeration_row = list(range(1, n_cols + 1))
      output_df.loc[-1] = numeration_row
      output_df = output_df.sort_index().reset_index(drop=True)
      output_df = output_df.astype(str)
      return output_df
